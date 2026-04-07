import csv
import io
import tempfile
from pathlib import Path

import PyPDF2
import docx
import openpyxl


def parse_pdf(file_bytes: bytes) -> str:
    reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text)
    return "\n\n".join(pages)


def parse_docx(file_bytes: bytes) -> str:
    doc = docx.Document(io.BytesIO(file_bytes))
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                paragraphs.append(row_text)

    return "\n\n".join(paragraphs)


def parse_xlsx(file_bytes: bytes) -> str:
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    sections = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            cell_values = [str(c) if c is not None else "" for c in row]
            if any(v.strip() for v in cell_values):
                rows.append(" | ".join(cell_values))
        if rows:
            sections.append(f"Sheet: {sheet_name}\n" + "\n".join(rows))

    wb.close()
    return "\n\n".join(sections)


def parse_csv_bytes(file_bytes: bytes) -> str:
    text = file_bytes.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(text))
    rows = []
    for row in reader:
        if any(cell.strip() for cell in row):
            rows.append(" | ".join(row))
    return "\n".join(rows)


def parse_text(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8", errors="replace")


PARSERS = {
    ".pdf": parse_pdf,
    ".docx": parse_docx,
    ".doc": parse_docx,
    ".xlsx": parse_xlsx,
    ".xls": parse_xlsx,
    ".csv": parse_csv_bytes,
    ".txt": parse_text,
    ".md": parse_text,
    ".json": parse_text,
}


def parse_file(filename: str, file_bytes: bytes) -> str:
    ext = Path(filename).suffix.lower()
    parser = PARSERS.get(ext)
    if parser is None:
        return parse_text(file_bytes)
    return parser(file_bytes)


def supported_extensions() -> list[str]:
    return list(PARSERS.keys())
